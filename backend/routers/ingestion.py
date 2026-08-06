"""Data ingestion routes — crawl + task status + history (DB-persisted).

Crawl runs asynchronously: POST returns immediately, frontend polls for status.
"""

from __future__ import annotations

import asyncio
import io
import uuid

from fastapi import APIRouter, BackgroundTasks, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from backend.auth import require_user
from backend.database import SessionLocal, get_db
from backend.models import CrawlTask, Competitor, User
from backend.schemas import (
    CrawlRequest,
    CrawlResponse,
    CrawlResponseData,
    CrawlTaskItem,
    CrawlTaskListResponse,
    TaskStatusData,
    TaskStatusResponse,
)
from backend.services.scraper import ScrapeError, fetch, truncate
from backend.services.vector_store import ingest

router = APIRouter(prefix="/api/v1")


def _run_crawl_sync(task_id: str, user_id: int, competitor_id: str, url: str, source_type: str):
    """Synchronous crawl logic — runs in a background thread."""
    db = SessionLocal()
    try:
        task = db.query(CrawlTask).filter(CrawlTask.task_id == task_id).first()
        if not task:
            return

        comp = db.query(Competitor).filter(
            Competitor.competitor_id == competitor_id,
            Competitor.user_id == user_id,
        ).first()

        # 1. Scrape (Playwright → httpx fallback)
        text = asyncio.run(fetch(url, timeout=25))
        text = truncate(text)

        # 2. Save content preview (first 3000 chars for UI display)
        preview = text[:3000]

        # 3. Ingest into vector store
        count = ingest(
            text=text,
            competitor_id=competitor_id,
            source_url=url,
            source_type=source_type,
        )

        # 3. Update competitor document_count
        if comp:
            comp.document_count = (comp.document_count or 0) + count

        # 5. Mark task completed
        task.status = "completed"
        task.progress_percentage = 100
        task.documents_created = count
        task.content_preview = preview
        db.commit()

    except ScrapeError as exc:
        if task:
            task.status = "failed"
            task.error_message = str(exc)
            db.commit()
    except Exception as exc:
        if task:
            task.status = "failed"
            task.error_message = f"采集异常: {exc}"
            db.commit()
    finally:
        db.close()


@router.post("/data/crawl", response_model=CrawlResponse)
def trigger_crawl(
    body: CrawlRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_user),
) -> CrawlResponse:
    """Start a crawl — returns immediately, runs scrape in background."""
    import uuid as _uuid
    task_id = f"crawl_{_uuid.uuid4().hex[:8]}"

    comp = db.query(Competitor).filter(
        Competitor.competitor_id == body.competitor_id,
        Competitor.user_id == current_user.id,
    ).first()
    comp_name = comp.name if comp else body.competitor_id

    # Create DB record
    db_task = CrawlTask(
        task_id=task_id,
        user_id=current_user.id,
        competitor_id=body.competitor_id,
        competitor_name=comp_name,
        source_url=body.url,
        source_type=body.source_type,
        status="processing",
    )
    db.add(db_task)
    db.commit()

    # Schedule background work
    background_tasks.add_task(
        _run_crawl_sync,
        task_id=task_id,
        user_id=current_user.id,
        competitor_id=body.competitor_id,
        url=body.url,
        source_type=body.source_type,
    )

    return CrawlResponse(
        data=CrawlResponseData(
            task_id=task_id,
            crawl_status="processing",
            estimated_time="30s",
        )
    )


@router.get("/data/task/{task_id}", response_model=TaskStatusResponse)
def get_task_status(
    task_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_user),
) -> TaskStatusResponse:
    """Get a single task's status (DB-backed)."""
    task = db.query(CrawlTask).filter(
        CrawlTask.task_id == task_id,
        CrawlTask.user_id == current_user.id,
    ).first()

    if task is None:
        return TaskStatusResponse(
            data=TaskStatusData(
                task_id=task_id,
                status="failed",
                progress_percentage=0,
                documents_created=0,
                error_message="任务记录不存在或已过期",
            )
        )

    return TaskStatusResponse(
        data=TaskStatusData(
            task_id=task.task_id,
            status=task.status,
            progress_percentage=task.progress_percentage,
            documents_created=task.documents_created,
            error_message=task.error_message,
            content_preview=task.content_preview,
        )
    )


@router.get("/data/tasks", response_model=CrawlTaskListResponse)
def list_crawl_tasks(
    db: Session = Depends(get_db),
    current_user: User = Depends(require_user),
) -> CrawlTaskListResponse:
    """List all crawl task history for the current user (newest first)."""
    tasks = (
        db.query(CrawlTask)
        .filter(CrawlTask.user_id == current_user.id)
        .order_by(CrawlTask.created_at.desc())
        .limit(50)
        .all()
    )

    return CrawlTaskListResponse(
        data=[
            CrawlTaskItem(
                task_id=t.task_id,
                competitor_id=t.competitor_id,
                competitor_name=t.competitor_name,
                source_url=t.source_url,
                source_type=t.source_type,
                status=t.status,
                progress_percentage=t.progress_percentage,
                documents_created=t.documents_created,
                error_message=t.error_message,
                content_preview=t.content_preview,
                created_at=t.created_at.isoformat() if t.created_at else "",
            )
            for t in tasks
        ]
    )


@router.get("/data/task/{task_id}/export")
def export_task_docx(
    task_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_user),
):
    """Export a completed task's scraped content as a .docx file."""
    task = db.query(CrawlTask).filter(
        CrawlTask.task_id == task_id,
        CrawlTask.user_id == current_user.id,
    ).first()

    if task is None:
        from fastapi.responses import JSONResponse
        return JSONResponse(
            status_code=404,
            content={"status": "error", "error_message": "任务记录不存在"},
        )

    if not task.content_preview:
        from fastapi.responses import JSONResponse
        return JSONResponse(
            status_code=400,
            content={"status": "error", "error_message": "该任务没有可供导出的内容"},
        )

    # Build .docx in memory
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import datetime

    doc = Document()

    # ---- Styles ----
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # ---- Title ----
    title = doc.add_heading(f"采集内容 · {task.competitor_name}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ---- Meta info ----
    meta = doc.add_paragraph()
    meta.style = doc.styles["Normal"]
    meta_run = meta.add_run(
        f"来源: {task.source_url}\n"
        f"采集时间: {task.created_at.strftime('%Y-%m-%d %H:%M:%S') if task.created_at else '—'}\n"
        f"入库片段数: {task.documents_created}\n"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = None

    doc.add_paragraph()  # spacer

    # ---- Content body ----
    for paragraph_text in task.content_preview.split("\n"):
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        if paragraph_text.strip():
            p.add_run(paragraph_text.strip())

    # ---- Write to buffer ----
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = task.competitor_name.replace(" ", "_").replace("/", "_")
    filename = f"{safe_name}_采集内容_{task.created_at.strftime('%Y%m%d') if task.created_at else 'export'}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
